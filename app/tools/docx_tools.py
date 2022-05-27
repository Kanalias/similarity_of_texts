from typing import List, Union
import xml.etree.ElementTree as ET
from docx import Document
from docx.text.paragraph import Paragraph


class DocExecuter:

    def __check_text(self, text):
        return text and not text.isspace()

    def get_paragraphs(self, doc: Document) -> List[Paragraph]:
        return [paragraph for paragraph in doc.paragraphs if self.__check_text(paragraph.text)]

    def get_texts(self, doc: Document, split_paragraphs: bool = True):
        texts = [paragraph.text for paragraph in doc.paragraphs if self.__check_text(paragraph.text)]

        if split_paragraphs:
            return texts

        return " ".join(texts)

    def __execute_text(self, elements, prefix: str):
        _text = ""

        for paragraph in elements.findall(f"{prefix}p"):
            for run in paragraph.findall(f"{prefix}r"):
                for text in run.findall(f"{prefix}t"):
                    _text += text.text

        return _text

    def execute_tables(self, doc: Document, split_cell: bool = True) -> List[str]:
        texts = []

        for table in doc.tables:
            xml_str = table._tbl.xml
            tree = ET.ElementTree(ET.fromstring(xml_str))
            root = tree.getroot()
            root_tag = root.tag
            prefix = root_tag[root_tag.find("{"): root_tag.rfind("}") + 1]
            rows = root.findall(f"{prefix}tr")

            for row in rows:
                _texts = []

                for cell in row.findall(f"{prefix}tc"):
                    text = self.__execute_text(cell, prefix)

                    if self.__check_text(text):
                        _texts.append(text)

                if _texts:
                    if split_cell:
                        texts.extend(_texts)
                    else:
                        texts.append(" ".join(_texts))

        return texts

    def execute_text(self, doc: Document,
                     split_cell: bool = True, split_paragraphs: bool = True,
                     only_text: bool = False) -> Union[List[str], str]:

        if only_text:
            split_paragraphs = False
            split_cell = False

        texts = self.get_texts(doc, split_paragraphs)
        text_tables = self.execute_tables(doc, split_cell=split_cell)

        if isinstance(texts, str):
            texts = [texts]

        texts.extend(text_tables)

        return " ".join(texts) if only_text else texts
